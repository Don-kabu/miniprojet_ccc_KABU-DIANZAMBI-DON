from docx import Document
from docx.shared import Pt, Inches,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from datetime import datetime
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from config import config


def add_footer(doc, left_text, center_text = datetime.strftime(datetime.now(),"%d/%m/%Y, %H:%M:%S"), right_text=""):
    section = doc.sections[0]
    footer = section.footer

    # Supprimer anciens paragraphes si besoin
    for p in footer.paragraphs:
        p.clear()

    # Créer un tableau 1x3
    table = footer.add_table(rows=1, cols=3, width=Inches(8.0))
    table.autofit = False  

    # Fixer les largeurs manuellement
    widths = [Inches(2.5), Inches(2.5), Inches(2.5)]
    for cell, width in zip(table.rows[0].cells, widths):
        cell.width = width

    cells = table.rows[0].cells

    # Texte gauche
    left_paragraph = cells[0].paragraphs[0]
    left_run = left_paragraph.add_run(left_text)
    left_run.font.size = Pt(10)

    # Texte centre
    center_paragraph = cells[1].paragraphs[0]
    center_run = center_paragraph.add_run(center_text)
    center_run.font.size = Pt(10)

    # Texte droite
    right_paragraph = cells[2].paragraphs[0]
    right_run = right_paragraph.add_run(right_text)
    right_run.font.size = Pt(10)






def set_cell_background(cell, color):
    """
    Applique une couleur de fond à une cellule
    :param cell: cellule du tableau (table.cell(row, col))
    :param color: code hex (ex: "FF0000" pour rouge)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)





# Création du document
doc = Document()
with open("out/RAW DATA 02/genarated.json","r") as file:
    data = json.load(file)


fill_doc_with_data(data)
# Sauvegarde du document
doc.save("table_personnalisee.docx")



