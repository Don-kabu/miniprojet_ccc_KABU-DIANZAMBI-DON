import sys
import os

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))


from docx import Document
from docx.shared import Pt, Inches,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import pandas as pd 
from config.config import candidate
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import qn

import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from config import config




def add_footer(doc, left_text, center_text = datetime.strftime(datetime.now(),"%d/%m/%Y, %H:%M:%S"), right_text=""):
    section = doc.sections[0]
    footer = section.footer

    # Supprimer anciens paragraphes si besoin
    for p in footer.paragraphs:
        p.clear()

    # Créer un tableau 1x3
    table = footer.add_table(rows=1, cols=3, width=Inches(8.0))
    table.autofit = False  

    # Fixer les largeurs manuellement
    widths = [Inches(2.5), Inches(2.5), Inches(2.5)]
    for cell, width in zip(table.rows[0].cells, widths):
        cell.width = width

    cells = table.rows[0].cells

    # Texte gauche
    left_paragraph = cells[0].paragraphs[0]
    left_run = left_paragraph.add_run(left_text)
    left_run.font.size = Pt(10)

    # Texte centre
    center_paragraph = cells[1].paragraphs[0]
    center_run = center_paragraph.add_run(center_text)
    center_run.font.size = Pt(10)

    # Texte droite
    right_paragraph = cells[2].paragraphs[0]
    right_run = right_paragraph.add_run(right_text)
    right_run.font.size = Pt(10)





def generate_csv_output(json_data:dict, output_path):
    # Préparation des données pour CSV
    csv_data = []
    for config in json_data:
        for measurement in config['measurements']:
            row = {
                'Sample': config['metadata']['sample'],
                # 'Configuration': config["parameters"]['name'],
                'RBW': config['parameters']['rbw'],
                'Frequency': measurement['frequency'],
                'Measurement': measurement['measurement'],
                'Limit': measurement['limit'],
                'Margin': measurement['margin'],
                'Verdict': measurement['verdict']
            }
            csv_data.append(row)
    
    # Écriture CSV
    df = pd.DataFrame(csv_data)
    df.to_csv(output_path, index=False)






















def fill_doc_with_data(input_data,doc,hash):
    headers = input_data[-1]["headers"]
    title=doc.add_heading("test repport".capitalize(),0)
    title.style.font.bold=True
    title.style.font.color.rgb = RGBColor(125,123,123)
    title.alignment =WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=0, cols=9)
    table.style = "Table Grid"  # Style avec bordures visibles
    table.style.font.color.rgb = RGBColor(0,0,0)
    header1 = table.add_row()
    for i in range(len(headers)):
        header1.cells[i].text = headers[i]
    try:
        for data in input_data[:-2]:
            
            row = table.add_row()
            row.cells[0].merge(row.cells[-1]).text=data["header"].split("\n")[0]


            row = table.add_row()
            row.cells[0].merge(row.cells[-1]).text=data["header"].split("\n")[1]

            if data["table"]:
                for measurement in data["table"]:
                    n_row=table.add_row()
                    for i in range(len(measurement)):
                        n_row.cells[i].text = str(measurement[i])
                        if measurement[i]=="FAIL":
                            n_row.cells[i].paragraphs[0].runs[0].font.bold =True
                            n_row.cells[i].paragraphs[0].runs[0].font.color.rgb =RGBColor(148, 16, 16)
                        elif measurement[i]=="PASS":
                            n_row.cells[i].paragraphs[0].runs[0].font.bold =True
                            n_row.cells[i].paragraphs[0].runs[0].font.color.rgb =RGBColor(0,128,0)
            else:
                row=table.add_row()
                row.cells[0].merge(row.cells[-1]).text="NO MEASUREMENT FOUND"
        try:
            row = table.add_row()
            row.cells[0].merge(row.cells[-2]).text="verdict"
            row.cells[-1].text=data["verdict"]
            if data["verdict"]=="FAIL":
                row.cells[i].paragraphs[0].runs[0].font.bold =True
                row.cells[i].paragraphs[0].runs[0].font.color.rgb =RGBColor(148, 16, 16)
            elif data["verdict"]=="PASS":
                row.cells[i].paragraphs[0].runs[0].font.bold =True
                row.cells[i].paragraphs[0].runs[0].font.color.rgb =RGBColor(0,128,0)
        except:
            pass
    except:
        pass

    add_footer(doc,config.candidate,right_text=hash)
